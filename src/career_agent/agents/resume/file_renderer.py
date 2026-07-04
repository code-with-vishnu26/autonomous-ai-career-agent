"""Real DOCX/PDF resume file generation (Phase 9, ADR-0033).

This is the first module that produces a binary artifact a company will
actually receive, so its layout rules are a locked spec, not styling
preferences (each one exists because real ATS parsers are known to choke
on its violation): single column, no tables, no text boxes, no images, no
headers/footers (contact info lives in the body), the five standard
section headings exactly ("Summary", "Work Experience", "Education",
"Skills", "Projects"), a standard font (Calibri 11pt), reverse-
chronological ordering. An empty section's heading is omitted entirely --
an empty "Projects" heading is noise to a parser and a human alike.

Content sourcing follows the same split the rest of this project enforces:

- **Tailored, gate-verified content** (summary, work highlights, skills,
  projects) comes from :class:`~career_agent.domain.models.TailoredContent`
  -- already through the truthfulness gate by the time this module runs
  (the pipeline only generates files for approved drafts), though this
  module independently re-verifies every ``source_entry_id`` it renders
  (via :func:`~career_agent.domain.rendering.resolve_work_entry`), raising
  ``KeyError`` loudly rather than silently dropping an entry -- the same
  "never trust that upstream already verified this" discipline as the
  plain-text renderer (ADR-0025).
- **Profile facts** (contact identity, real employment dates, company
  names, and the entire Education section) come read-only from
  :class:`~career_agent.domain.models.MasterProfile`. Education
  deliberately does NOT pass through ``TailoredContent`` at all
  (ADR-0033): it is a verified profile fact, not content to tailor --
  adding it to the generated-content path would expand the truthfulness
  gate's surface for zero benefit. The generator structurally cannot
  write, reorder, or fabricate education, because no generated type
  carries it.

**Determinism (ADR-0033):** a DOCX is a zip archive, and python-docx
stamps each zip entry with the current wall-clock second -- so two
byte-identical renders straddling a second boundary would hash
differently, breaking the content-addressed never-overwrite scheme.
``_normalize_zip_timestamps`` rewrites the archive with a fixed epoch and
sorted entries, making the DOCX bytes a pure function of
(content, profile). Verified empirically in this environment: raw
python-docx output is NOT cross-second deterministic; normalized output
is, and still opens as a valid document (and still converts via
LibreOffice). The PDF is a *derived view* and is NOT byte-deterministic
(LibreOffice embeds a CreationDate) -- its hash still makes it traceable
and collision-safe, but the DOCX is the canonical, reproducible artifact.

**Never overwrite, by construction:** the content hash is embedded in the
filename (``resume-{resume_id}-{hash12}.docx``). Regenerating after any
content or profile change produces a different hash and therefore a
different file; regenerating with identical inputs produces the identical
file (idempotent). No overwrite check is needed because no two different
contents can ever share a name.

**PDF conversion is environment-dependent, checked at runtime, never
assumed:** this sandbox shipped ``soffice`` with only ``libreoffice-core``
-- the binary existed but could not load a DOCX until
``libreoffice-writer`` was installed, a real failure mode discovered
empirically, not hypothesized. A missing or failing converter raises the
typed :class:`PdfConversionUnavailableError`; the DOCX (which is what
Lever's upload accepts) is already on disk by then. ReportLab remains a
named, unbuilt fallback -- only if a real environment turns out to lack
LibreOffice entirely (ADR-0033's revisit criteria).
"""

from __future__ import annotations

import hashlib
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

from career_agent.domain.models import (
    EducationEntry,
    MasterProfile,
    ResumeArtifact,
    TailoredContent,
)
from career_agent.domain.rendering import resolve_work_entry

_FIXED_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
_HASH_PREFIX_LEN = 12


class PdfConversionUnavailableError(Exception):
    """LibreOffice (soffice) is missing or failed to convert the DOCX.

    Typed and named so a caller can tell "the PDF view could not be
    produced in this environment" apart from any content problem -- the
    DOCX artifact already exists on disk when this is raised, so nothing
    about the application itself is blocked, only the optional PDF view.
    """


def render_resume_docx(
    resume_id: str,
    content: TailoredContent,
    profile: MasterProfile,
    artifacts_dir: Path,
) -> ResumeArtifact:
    """Render one deterministic, ATS-safe DOCX and return its artifact record.

    Raises ``KeyError`` if any work/project ``source_entry_id`` cannot be
    resolved against ``profile`` -- independently re-verified here, never
    assumed already checked upstream.
    """
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    document = _build_document(content, profile)
    # Checked AFTER all content is added, immediately before save -- checking
    # the fresh document would miss anything added during building, a gap the
    # Phase 9 injection pass found empirically (ADR-0033).
    _forbid_unsafe_constructs(document)

    raw_path = artifacts_dir / f".building-{resume_id}.docx"
    document.save(str(raw_path))
    normalized_path = artifacts_dir / f".normalized-{resume_id}.docx"
    _normalize_zip_timestamps(raw_path, normalized_path)
    raw_path.unlink()

    content_hash = _sha256(normalized_path)
    final_path = artifacts_dir / (
        f"resume-{resume_id}-{content_hash[:_HASH_PREFIX_LEN]}.docx"
    )
    # Identical content produces the identical name (idempotent); different
    # content can never collide with an existing name, so replace() here can
    # only ever overwrite a byte-identical file.
    normalized_path.replace(final_path)

    return ResumeArtifact(
        resume_id=resume_id,
        profile_version=profile.version,
        format="docx",
        path=str(final_path),
        content_hash=content_hash,
    )


def convert_to_pdf(
    docx_artifact: ResumeArtifact,
    artifacts_dir: Path,
    *,
    soffice_executable: str = "soffice",
) -> ResumeArtifact:
    """Convert an existing DOCX artifact to a text-based PDF via LibreOffice.

    Raises :class:`PdfConversionUnavailableError` when the converter is
    missing or fails -- never silently returns without a PDF. The PDF's
    bytes are not reproducible run-to-run (LibreOffice embeds a creation
    timestamp); its hash is still recorded for traceability.
    """
    if shutil.which(soffice_executable) is None:
        raise PdfConversionUnavailableError(
            f"{soffice_executable!r} not found on PATH -- cannot produce a "
            f"PDF in this environment (the DOCX artifact at "
            f"{docx_artifact.path!r} already exists and is unaffected)"
        )
    docx_path = Path(docx_artifact.path)
    result = subprocess.run(  # noqa: S603 -- fixed argv, no shell
        [
            soffice_executable,
            "--headless",
            "--convert-to",
            "pdf",
            str(docx_path),
            "--outdir",
            str(artifacts_dir),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )
    produced = artifacts_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not produced.exists():
        detail = result.stderr.strip() or result.stdout.strip()
        raise PdfConversionUnavailableError(
            f"soffice failed to convert {docx_path.name!r} "
            f"(exit={result.returncode}): {detail}"
        )

    content_hash = _sha256(produced)
    final_path = artifacts_dir / (
        f"resume-{docx_artifact.resume_id}-{content_hash[:_HASH_PREFIX_LEN]}.pdf"
    )
    produced.replace(final_path)
    return ResumeArtifact(
        resume_id=docx_artifact.resume_id,
        profile_version=docx_artifact.profile_version,
        format="pdf",
        path=str(final_path),
        content_hash=content_hash,
    )


def _build_document(content: TailoredContent, profile: MasterProfile) -> Document:
    document = Document()
    _apply_base_font(document)

    # Contact identity, in the body -- never a header/footer, which many
    # ATS parsers skip entirely.
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(profile.basics.name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    contact_bits = [profile.basics.email]
    if profile.basics.phone:
        contact_bits.append(profile.basics.phone)
    if profile.basics.location:
        contact_bits.append(profile.basics.location)
    document.add_paragraph(" | ".join(contact_bits))

    document.add_heading("Summary", level=1)
    document.add_paragraph(content.summary)

    if content.work:
        document.add_heading("Work Experience", level=1)
        entries = [
            (entry, resolve_work_entry(entry, profile)) for entry in content.work
        ]
        entries.sort(key=lambda pair: pair[1].start_date, reverse=True)
        for tailored, source in entries:
            end_label = (
                source.end_date.isoformat() if source.end_date else "Present"
            )
            line = document.add_paragraph()
            run = line.add_run(f"{tailored.position} — {source.name}")
            run.bold = True
            document.add_paragraph(f"{source.start_date.isoformat()} - {end_label}")
            for highlight in tailored.highlights:
                document.add_paragraph(highlight, style="List Bullet")

    education = sorted(
        profile.education,
        key=lambda entry: (entry.start_date is not None, entry.start_date),
        reverse=True,
    )
    if education:
        document.add_heading("Education", level=1)
        for entry in education:
            document.add_paragraph(_format_education(entry))

    if content.skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(content.skills))

    if content.projects:
        document.add_heading("Projects", level=1)
        project_ids = {project.id for project in profile.projects}
        for project in content.projects:
            if project.source_entry_id not in project_ids:
                raise KeyError(
                    f"no ProjectEntry with id={project.source_entry_id!r} in "
                    f"profile version={profile.version!r}"
                )
            line = document.add_paragraph()
            line.add_run(project.name).bold = True
            for highlight in project.highlights:
                document.add_paragraph(highlight, style="List Bullet")

    return document


def _format_education(entry: EducationEntry) -> str:
    """One education line, rendered verbatim from profile facts (ADR-0033).

    Nothing here is generated, reworded, or reordered semantically -- the
    ordering applied by the caller is plain reverse-chronology, and every
    string is the profile's own.
    """
    degree_bits = [bit for bit in (entry.study_type, entry.area) if bit]
    degree = ", ".join(degree_bits)
    label = f"{degree} — {entry.institution}" if degree else entry.institution
    if entry.start_date:
        end_label = entry.end_date.isoformat() if entry.end_date else "Present"
        return f"{label} ({entry.start_date.isoformat()} - {end_label})"
    return label


def _apply_base_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _forbid_unsafe_constructs(document: Document) -> None:
    """Assert the ATS-unsafe constructs are absent from the FINISHED document.

    Runs immediately before save, after all content is added -- the Phase 9
    injection pass proved that checking the fresh document instead would
    miss a table introduced during building (the injected violation slipped
    past a fresh-document check and was only caught by the file-reading
    test). If a future code change ever adds a table or an image anywhere
    in the build path, generation fails loudly here instead of shipping a
    file real parsers would mangle.
    """
    if document.tables:
        raise ValueError("ATS-safe layout spec forbids tables (ADR-0033)")
    if document.inline_shapes:
        raise ValueError("ATS-safe layout spec forbids images (ADR-0033)")


def _normalize_zip_timestamps(source: Path, destination: Path) -> None:
    """Rewrite a zip (DOCX) with fixed timestamps and sorted entries.

    python-docx stamps zip entries with the current wall-clock second, so
    otherwise-identical renders straddling a second boundary would hash
    differently. Fixing the timestamp and entry order makes the bytes a
    pure function of the document content -- verified to keep the file a
    valid, LibreOffice-convertible DOCX.
    """
    with (
        zipfile.ZipFile(source) as zin,
        zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        for item in sorted(zin.infolist(), key=lambda info: info.filename):
            normalized = zipfile.ZipInfo(
                item.filename, date_time=_FIXED_ZIP_TIMESTAMP
            )
            normalized.compress_type = zipfile.ZIP_DEFLATED
            normalized.external_attr = item.external_attr
            zout.writestr(normalized, zin.read(item.filename))


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()
