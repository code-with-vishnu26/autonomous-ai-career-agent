"""Phase 9 / ADR-0033: real DOCX/PDF resume file generation.

Every layout assertion here reads the generated DOCX back through
python-docx against the real file on disk -- never trusts that the writer
"should have" produced the spec. Education is proven to come read-only
from MasterProfile (the authoritative Option (a) decision): rendered
verbatim, unmutated, structurally impossible for tailored content to
override because no generated type carries an education field at all.
"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from career_agent.agents.resume.file_renderer import (
    PdfConversionUnavailableError,
    convert_to_pdf,
    render_resume_docx,
)
from career_agent.domain.models import (
    DraftedTailoring,
    EducationEntry,
    MasterProfile,
    TailoredContent,
    TailoredProjectEntry,
    TailoredWorkEntry,
)
from tests.agents._profile_fixture import sample_master_profile


def _profile_with_education(entries: list[EducationEntry]) -> MasterProfile:
    return sample_master_profile().model_copy(update={"education": entries})


_EDUCATION = [
    EducationEntry(
        id="edu-mit",
        institution="MIT",
        area="Computer Science",
        study_type="Master's",
        start_date=date(2019, 9, 1),
        end_date=date(2021, 6, 1),
    ),
    EducationEntry(
        id="edu-state",
        institution="State University",
        area="Mathematics",
        study_type="Bachelor's",
        start_date=date(2015, 9, 1),
        end_date=date(2019, 6, 1),
    ),
]


def _content() -> TailoredContent:
    return TailoredContent(
        summary="Engineer focused on reliable systems.",
        work=[
            TailoredWorkEntry(
                source_entry_id="work-techco",
                position="Software Engineer",
                highlights=["Built REST APIs serving 2M requests/day"],
            )
        ],
        skills=["Python", "Django"],
        projects=[
            TailoredProjectEntry(
                source_entry_id="proj-internal",
                name="Internal Tool",
                highlights=["Built an internal tool"],
            )
        ],
    )


def _all_text(path: str) -> str:
    return "\n".join(p.text for p in Document(path).paragraphs)


def _headings(path: str) -> list[str]:
    return [
        p.text
        for p in Document(path).paragraphs
        if p.style.name.startswith("Heading")
    ]


# ---------------------------------------------------------------------------
# ATS-safe layout spec, verified against the real generated file
# ---------------------------------------------------------------------------


def test_standard_section_headings_exactly(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    assert _headings(artifact.path) == [
        "Summary",
        "Work Experience",
        "Education",
        "Skills",
        "Projects",
    ]


def test_no_tables_images_or_header_footer_content(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    document = Document(artifact.path)
    assert len(document.tables) == 0
    assert len(document.inline_shapes) == 0
    for section in document.sections:
        header_text = "".join(p.text for p in section.header.paragraphs).strip()
        footer_text = "".join(p.text for p in section.footer.paragraphs).strip()
        assert header_text == ""
        assert footer_text == ""


def test_contact_info_is_in_the_body_not_a_header(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "Ada Lovelace" in body
    assert "ada@example.com" in body


def test_no_links_configured_renders_no_extra_line(tmp_path: Path) -> None:
    """Phase 72/ADR-0090: an existing profile with none of the new link
    fields set renders identically to before this phase -- no blank line,
    no empty second contact line."""
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    paragraphs = [p.text for p in Document(artifact.path).paragraphs]
    contact_index = paragraphs.index("ada@example.com")
    assert paragraphs[contact_index + 1] == "Summary"


def test_profile_links_render_on_a_second_contact_line(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    profile = profile.model_copy(
        update={
            "basics": profile.basics.model_copy(
                update={
                    "linkedin_url": "https://linkedin.com/in/ada",
                    "github_url": "https://github.com/ada",
                    "website_url": "https://ada.dev",
                    "other_links": ["https://ada.dev/talks"],
                }
            )
        }
    )
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "https://linkedin.com/in/ada" in body
    assert "https://github.com/ada" in body
    assert "https://ada.dev" in body
    assert "https://ada.dev/talks" in body


def test_project_url_renders_next_to_the_project_name(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    profile = profile.model_copy(
        update={
            "projects": [
                project.model_copy(update={"url": "https://github.com/ada/tool"})
                for project in profile.projects
            ]
        }
    )
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "https://github.com/ada/tool" in body


def test_project_without_a_url_renders_no_parenthetical(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "()" not in body


def test_base_font_is_calibri_11(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    style = Document(artifact.path).styles["Normal"]
    assert style.font.name == "Calibri"
    assert style.font.size.pt == 11


def test_empty_sections_omit_their_headings(tmp_path: Path) -> None:
    """An empty section's heading is noise to a parser -- omitted entirely."""
    profile = _profile_with_education([])  # no education
    content = TailoredContent(summary="Just a summary.")  # no work/skills/projects
    artifact = render_resume_docx("r-1", content, profile, tmp_path)
    assert _headings(artifact.path) == ["Summary"]


def test_real_dates_and_company_names_from_the_profile(tmp_path: Path) -> None:
    """The generator structurally can't write dates/companies; the profile's
    real ones must still reach the file -- via resolve_work_entry."""
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "Techco" in body
    assert "2022-01-01 - Present" in body


def test_fabricated_source_entry_id_raises_loudly(tmp_path: Path) -> None:
    """Same canary as the plain-text renderer: an unresolvable reference is
    a loud KeyError, never a silently dropped entry."""
    profile = _profile_with_education(_EDUCATION)
    content = _content().model_copy(
        update={
            "work": [
                TailoredWorkEntry(
                    source_entry_id="work-nonexistent",
                    position="Staff Engineer",
                    highlights=["x"],
                )
            ]
        }
    )
    with pytest.raises(KeyError, match="work-nonexistent"):
        render_resume_docx("r-1", content, profile, tmp_path)


# ---------------------------------------------------------------------------
# Education: read-only from MasterProfile (Option (a), ADR-0033)
# ---------------------------------------------------------------------------


def test_education_is_rendered_from_master_profile_verbatim(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "Master's, Computer Science — MIT (2019-09-01 - 2021-06-01)" in body


def test_multiple_education_entries_preserved_reverse_chronological(
    tmp_path: Path,
) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    body = _all_text(artifact.path)
    assert "MIT" in body
    assert "State University" in body
    assert body.index("MIT") < body.index("State University")  # newest first


def test_missing_education_is_handled_safely(tmp_path: Path) -> None:
    profile = _profile_with_education([])
    artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    assert "Education" not in _headings(artifact.path)


def test_tailored_content_structurally_cannot_carry_education() -> None:
    """The Option (a) guarantee at its strongest: no generated type has an
    education field at all, so no drafter/generator output can ever
    override, reorder, or fabricate education -- there is nowhere to put
    it. Structural, not a runtime check."""
    assert "education" not in TailoredContent.model_fields
    assert "education" not in DraftedTailoring.model_fields


def test_rendering_does_not_mutate_the_profile(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    before = profile.model_dump(mode="json")
    render_resume_docx("r-1", _content(), profile, tmp_path)
    assert profile.model_dump(mode="json") == before


# ---------------------------------------------------------------------------
# Determinism + never-overwrite (content-addressed filenames)
# ---------------------------------------------------------------------------


def test_identical_inputs_produce_byte_identical_docx(tmp_path: Path) -> None:
    """python-docx alone is NOT cross-second deterministic (zip timestamps)
    -- verified empirically; the normalization step makes the bytes a pure
    function of (content, profile).

    The sleep is 2.1s, not ~1s: ZIP DOS timestamps have TWO-second
    granularity, so a shorter sleep can land both renders in the same
    2-second bucket and let an un-normalized file pass by luck -- exactly
    what happened on the first injection attempt against this test
    (normalization removed, test still green). 2.1s guarantees the raw
    timestamps differ, so only genuine normalization can make the hashes
    match."""
    import time

    profile = _profile_with_education(_EDUCATION)
    first = render_resume_docx("r-1", _content(), profile, tmp_path)
    time.sleep(2.1)  # guarantee crossing a 2s DOS-timestamp boundary
    second = render_resume_docx("r-1", _content(), profile, tmp_path)
    assert first.content_hash == second.content_hash
    assert first.path == second.path  # idempotent, not a duplicate file


def test_changed_content_gets_a_new_file_and_never_touches_the_old_one(
    tmp_path: Path,
) -> None:
    profile = _profile_with_education(_EDUCATION)
    first = render_resume_docx("r-1", _content(), profile, tmp_path)
    original_bytes = Path(first.path).read_bytes()

    changed = _content().model_copy(update={"summary": "A different summary."})
    second = render_resume_docx("r-1", changed, profile, tmp_path)

    assert second.content_hash != first.content_hash
    assert second.path != first.path
    assert Path(first.path).read_bytes() == original_bytes  # untouched
    assert Path(second.path).exists()


def test_artifact_records_trace_back_to_resume_and_profile_version(
    tmp_path: Path,
) -> None:
    profile = _profile_with_education(_EDUCATION)
    artifact = render_resume_docx("resume-abc", _content(), profile, tmp_path)
    assert artifact.resume_id == "resume-abc"
    assert artifact.profile_version == "profile-v1"
    assert artifact.format == "docx"
    assert artifact.content_hash[:12] in artifact.path


# ---------------------------------------------------------------------------
# PDF conversion: real when the environment has it, typed refusal when not
# ---------------------------------------------------------------------------


def test_missing_soffice_raises_typed_error_and_docx_survives(
    tmp_path: Path,
) -> None:
    profile = _profile_with_education(_EDUCATION)
    docx_artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    with pytest.raises(PdfConversionUnavailableError, match="not found on PATH"):
        convert_to_pdf(
            docx_artifact, tmp_path, soffice_executable="soffice-does-not-exist"
        )
    assert Path(docx_artifact.path).exists()  # the DOCX is unaffected


@pytest.mark.skipif(
    shutil.which("soffice") is None,
    reason="LibreOffice (soffice) not available in this environment",
)
def test_real_pdf_conversion_produces_a_text_based_pdf(tmp_path: Path) -> None:
    profile = _profile_with_education(_EDUCATION)
    docx_artifact = render_resume_docx("r-1", _content(), profile, tmp_path)
    try:
        pdf_artifact = convert_to_pdf(docx_artifact, tmp_path)
    except PdfConversionUnavailableError as exc:
        # soffice can exist without libreoffice-writer (this sandbox's own
        # out-of-the-box state, ADR-0033) -- an environment gap, not a code
        # failure; the typed error firing is itself correct behavior.
        pytest.skip(f"soffice present but cannot convert here: {exc}")

    data = Path(pdf_artifact.path).read_bytes()
    assert data[:5] == b"%PDF-"
    assert b"/Font" in data  # real text with font resources, not an image render
    assert pdf_artifact.format == "pdf"
    assert pdf_artifact.resume_id == docx_artifact.resume_id
    assert pdf_artifact.profile_version == docx_artifact.profile_version
